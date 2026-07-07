"""
ocr_engine.py
=============
Turns an uploaded file (PDF / PNG / JPG / DOCX / DOC / TXT) into
plain text, regardless of format.

DESIGN DECISION - lazy-loaded OCR:
EasyOCR depends on PyTorch and downloads a ~50-100MB neural network
model the FIRST time it runs. That is too slow/heavy to load every
time this module is imported (e.g. every FastAPI worker startup, or
every `pytest` run). So we load the EasyOCR reader lazily - only the
first time it's actually needed - and cache it afterward.

STRATEGY PER FILE TYPE:
  - .txt          -> just read the file directly, no OCR needed
  - .docx / .doc  -> use python-docx to pull paragraph + table text
  - .pdf          -> try PyMuPDF's native text extraction first
                     (fast, exact, works for digitally-generated PDFs).
                     If that returns near-empty text (i.e. the PDF is
                     just scanned images), fall back to rendering each
                     page as an image and running EasyOCR on it.
                     Tables are extracted separately with pdfplumber.
  - .png/.jpg/...  -> run EasyOCR directly on the image

Every path returns a tuple: (extracted_text, confidence_score, tables)
so the caller always gets a consistent shape back.
"""

import logging
from pathlib import Path
from typing import List, Optional, Tuple

from utils import clean_ocr_text

logger = logging.getLogger(__name__)

# ------------------------------------------------------------
# Lazy-loaded EasyOCR reader (expensive to initialize)
# ------------------------------------------------------------
_easyocr_reader = None


def _get_easyocr_reader():
    """
    Creates the EasyOCR Reader object once and reuses it for every
    subsequent call. `gpu=False` because we can't assume the
    deployment machine has a GPU - EasyOCR runs fine on CPU, just
    somewhat slower.
    """
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr  # imported here, not at module top, to keep
                         # startup fast when OCR isn't needed yet
        logger.info("Loading EasyOCR model (first use only, may take a moment)...")
        _easyocr_reader = easyocr.Reader(["en"], gpu=False)
    return _easyocr_reader


def _ocr_image_file(path: Path) -> Tuple[str, float]:
    """Runs EasyOCR on a single image file. Returns (text, avg_confidence)."""
    reader = _get_easyocr_reader()
    results = reader.readtext(str(path))  # list of (bbox, text, confidence)
    if not results:
        return "", 0.0
    lines = [text for (_bbox, text, _conf) in results]
    confidences = [conf for (_bbox, _text, conf) in results]
    avg_confidence = sum(confidences) / len(confidences)
    return "\n".join(lines), avg_confidence


def _extract_pdf_native_text(path: Path) -> str:
    """
    Uses PyMuPDF (imported as `fitz`) to pull text directly out of a
    PDF's internal structure. This works instantly and perfectly for
    PDFs that were generated digitally (e.g. exported from a hospital
    LIS system) - no image recognition needed at all.
    """
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def _extract_pdf_via_ocr(path: Path) -> Tuple[str, float]:
    """
    Fallback for scanned PDFs: renders each page to an image using
    PyMuPDF, then runs EasyOCR on each rendered page image.
    """
    import fitz
    doc = fitz.open(str(path))
    all_text = []
    all_confidences = []

    for page_index in range(len(doc)):
        page = doc[page_index]
        # Render at 2x zoom for better OCR accuracy on small fonts
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        temp_image_path = path.parent / f"_ocr_page_{page_index}.png"
        pix.save(str(temp_image_path))
        try:
            text, confidence = _ocr_image_file(temp_image_path)
            all_text.append(text)
            all_confidences.append(confidence)
        finally:
            temp_image_path.unlink(missing_ok=True)  # always clean up temp file

    doc.close()
    combined_text = "\n".join(all_text)
    avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0.0
    return combined_text, avg_confidence


def _extract_pdf_tables(path: Path) -> List[List[List[Optional[str]]]]:
    """
    Uses pdfplumber to extract tables as nested lists (list of tables,
    each table is a list of rows, each row is a list of cell strings).
    Many lab reports format test results as tables, so this gives
    parser.py a structured alternative to plain text when available.
    Returns an empty list if no tables are found or the PDF is scanned
    images (pdfplumber can't read table structure from pixels).
    """
    import pdfplumber
    tables: List[List[List[Optional[str]]]] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    if table:
                        tables.append(table)
    except Exception as exc:  # pdfplumber can fail on some malformed PDFs
        logger.warning("pdfplumber table extraction failed for %s: %s", path.name, exc)
    return tables


def _extract_docx_text(path: Path) -> str:
    """Extracts paragraph text AND table cell text from a .docx file."""
    import docx  # python-docx
    document = docx.Document(str(path))

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


# ------------------------------------------------------------
# Public entry point
# ------------------------------------------------------------
def extract_text_from_file(path: Path, file_ext: str) -> Tuple[str, float, list]:
    """
    Main function called by routes.py after a file is uploaded.

    Returns:
        raw_text (str)       - cleaned extracted text
        confidence (float)   - 0.0-1.0. 1.0 for text formats where
                                extraction is exact (txt/docx/native pdf),
                                the EasyOCR average confidence otherwise.
        tables (list)        - extracted tables, if any (PDF only for now)
    """
    file_ext = file_ext.lower()
    tables: list = []

    try:
        if file_ext == ".txt":
            raw_text = path.read_text(encoding="utf-8", errors="ignore")
            confidence = 1.0

        elif file_ext in (".docx", ".doc"):
            raw_text = _extract_docx_text(path)
            confidence = 1.0

        elif file_ext == ".pdf":
            native_text = _extract_pdf_native_text(path)
            tables = _extract_pdf_tables(path)

            # Heuristic: if native extraction found almost no text,
            # this PDF is very likely scanned images, not real text.
            if len(native_text.strip()) < 30:
                logger.info("PDF '%s' looks scanned - falling back to OCR", path.name)
                raw_text, confidence = _extract_pdf_via_ocr(path)
            else:
                raw_text, confidence = native_text, 1.0

        elif file_ext in (".png", ".jpg", ".jpeg"):
            raw_text, confidence = _ocr_image_file(path)

        else:
            raise ValueError(f"Unsupported file extension for OCR: {file_ext}")

    except Exception as exc:
        logger.error("Text extraction failed for %s: %s", path.name, exc)
        raise

    cleaned_text = clean_ocr_text(raw_text)
    return cleaned_text, round(float(confidence), 3), tables
